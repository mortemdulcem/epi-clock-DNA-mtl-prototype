"""
Robust manuscript translator with checkpointing.
- Saves the .docx every N paragraphs
- Caches translations on disk so re-runs skip done work
- Handles rate limits with exponential backoff
- Preserves all 31 tables, 15 figures, formatting, references
"""

import os
import re
import json
import time
import signal
import sys
from docx import Document
from deep_translator import GoogleTranslator

SRC = "attached_assets/DNA_Metilasyon_Saatleriyle_Bağımlılık_1777114079284.docx"
OUT_DIR = "figures/output/submission_package_EN"
OUT = os.path.join(OUT_DIR, "02_Manuscript_Anonymized.docx")
CACHE = "/tmp/translate_cache.json"
CHECKPOINT = "/tmp/translate_checkpoint.docx"
os.makedirs(OUT_DIR, exist_ok=True)

# ---- glossary ----
GLOSSARY = {
    "GrimAge": "GRIMAGE_KEEP",
    "PhenoAge": "PHENOAGE_KEEP",
    "DunedinPACE": "DUNEDINPACE_KEEP",
    "epi-clock-prototype": "EPICLOCKPROTOTYPE_KEEP",
    "Adli Tıp Anabilim Dalı": "Department of Forensic Medicine",
    "Ankara Bilkent Şehir Hastanesi": "Ankara Bilkent City Hospital",
    "epigenetik yaş ivmelenmesi": "epigenetic age acceleration",
    "epigenetik yaş hızlanması": "epigenetic age acceleration",
    "madde kullanım bozukluğu": "substance use disorder",
    "madde kullanım bozuklukları": "substance use disorders",
    "alkol kullanım bozukluğu": "alcohol use disorder",
    "opioid kullanım bozukluğu": "opioid use disorder",
}
ANON_REPLACEMENTS = {
    "Nurcan Denli Bayır, MD¹*": "[Author]¹*",
    "Nurcan Denli Bayır, MD": "[Author]",
    "Nurcan Denli Bayır": "[Author]",
    "Nurcan Bayır": "[Author]",
    "drnurcandenlibayir@gmail.com": "[email redacted]",
    "0009-00042874-4594": "[ORCID redacted]",
    "0009-0004-2874-4594": "[ORCID redacted]",
    "Adli Tıp Anabilim Dalı, Ankara Bilkent Şehir Hastanesi, Ankara, Türkiye":
        "[Department, Institution, Country]",
    "Department of Forensic Medicine, Ankara Bilkent City Hospital, Ankara, Türkiye":
        "[Department, Institution, Country]",
    "Ankara Bilkent Şehir Hastanesi": "[Institution]",
    "Ankara Bilkent City Hospital": "[Institution]",
}

TR_CHARS = set("çğıöşüÇĞİÖŞÜâîû")


def looks_turkish(text):
    if not text or not text.strip():
        return False
    s = text.strip()
    if any(c in TR_CHARS for c in s):
        return True
    if len(s) < 6:
        return False
    if re.match(r"^[\x00-\x7f\s\d\W]+$", s):
        markers = re.findall(
            r"\b(ve|bu|bir|ile|olarak|için|olan|gore|göre|"
            r"çalışma|sonuç|yöntem|tartışma|"
            r"madde|bağımlılık|metilasyon|epigenetik|yaş|hastalık)\b",
            s.lower(),
        )
        return len(markers) >= 2
    return True


def apply_glossary(text):
    out = text
    for tr, en in GLOSSARY.items():
        out = re.sub(re.escape(tr), en, out, flags=re.IGNORECASE)
    return out


def restore_keep_tokens(text):
    return (text
            .replace("EPICLOCKPROTOTYPE_KEEP", "epi-clock-prototype")
            .replace("GRIMAGE_KEEP", "GrimAge")
            .replace("PHENOAGE_KEEP", "PhenoAge")
            .replace("DUNEDINPACE_KEEP", "DunedinPACE"))


def anonymize(text):
    out = text
    for k, v in ANON_REPLACEMENTS.items():
        out = out.replace(k, v)
    return out


# ---- cache ----
if os.path.exists(CACHE):
    with open(CACHE) as f:
        _cache = json.load(f)
    print(f"[+] loaded cache with {len(_cache)} entries")
else:
    _cache = {}


def save_cache():
    with open(CACHE, "w") as f:
        json.dump(_cache, f, ensure_ascii=False)


_translator = GoogleTranslator(source="tr", target="en")


def translate(text):
    if not text or not text.strip():
        return text
    key = text
    if key in _cache:
        return _cache[key]

    if len(text) > 4500:
        # split by sentence
        chunks = re.split(r"(?<=[\.!?])\s+", text)
        out, buf = [], ""
        for c in chunks:
            if len(buf) + len(c) > 4000:
                out.append(translate(buf))
                buf = c
            else:
                buf = (buf + " " + c).strip()
        if buf:
            out.append(translate(buf))
        result = " ".join(out)
        _cache[key] = result
        return result

    glossed = apply_glossary(text)
    last_err = None
    for attempt in range(5):
        try:
            r = _translator.translate(glossed)
            if r:
                r = restore_keep_tokens(r)
                _cache[key] = r
                return r
        except Exception as e:
            last_err = e
            wait = 2 ** attempt
            print(f"   ! retry {attempt+1}/5 after {wait}s ({type(e).__name__}: {str(e)[:60]})")
            time.sleep(wait)
    print(f"   ! FAILED after 5 retries; keeping Turkish: {text[:60]!r}")
    return text


def replace_paragraph_text(p, new_text):
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


def process_paragraph(p):
    full = p.text
    if not full.strip():
        return False
    if not looks_turkish(full):
        new = anonymize(full)
        if new != full:
            replace_paragraph_text(p, new)
            return True
        return False
    new = translate(full)
    new = anonymize(new)
    if new != full:
        replace_paragraph_text(p, new)
        return True
    return False


def main():
    print("[+] Loading source...")
    if os.path.exists(CHECKPOINT):
        print("[+] resuming from checkpoint")
        doc = Document(CHECKPOINT)
    else:
        doc = Document(SRC)

    paragraphs = list(doc.paragraphs)
    n_par = len(paragraphs)
    n_tab = len(doc.tables)
    print(f"[+] paragraphs={n_par}  tables={n_tab}  images={len(doc.inline_shapes)}")

    save_every = 30
    t0 = time.time()
    for i, p in enumerate(paragraphs):
        process_paragraph(p)
        if (i + 1) % save_every == 0:
            doc.save(CHECKPOINT)
            save_cache()
            elapsed = time.time() - t0
            rate = (i + 1) / elapsed
            eta = (n_par - i - 1) / rate if rate else 0
            print(f"   [{i+1}/{n_par}] checkpoint saved | {rate:.1f} par/s | ETA {eta:.0f}s")
            sys.stdout.flush()

    doc.save(CHECKPOINT)
    save_cache()
    print(f"[+] body done in {time.time()-t0:.1f}s")

    print(f"[+] translating {n_tab} tables...")
    t0 = time.time()
    for ti, t in enumerate(doc.tables):
        for row in t.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    process_paragraph(cp)
        if (ti + 1) % 5 == 0:
            doc.save(CHECKPOINT)
            save_cache()
            print(f"   [table {ti+1}/{n_tab}] checkpoint saved")
            sys.stdout.flush()

    # banner
    banner = doc.paragraphs[0].insert_paragraph_before(
        "[Author identifying information removed for blinded peer review]"
    )
    for r in banner.runs:
        r.italic = True

    doc.save(OUT)
    save_cache()
    print(f"[+] tables done in {time.time()-t0:.1f}s")
    print(f"[+] FINAL saved: {OUT}")


if __name__ == "__main__":
    main()
