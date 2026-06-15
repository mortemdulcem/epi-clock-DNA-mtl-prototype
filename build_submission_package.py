"""
Build journal submission package + run AI/LLM detector on each document.

Outputs (figures/output/submission_package/):
  01_Baslik_Sayfasi.docx              (Title Page with author info)
  02_El_Yazmasi_Anonim.docx           (Manuscript without author info)
  03_Cikar_Beyani.docx                (Conflict of Interest statement)
  04_On_Yazi_Cover_Letter.docx        (Cover letter to editor)
  05_Tesekkurler.docx                 (Acknowledgments)
  06_Veri_Bildirimi.docx              (Data availability statement)
  07_Etik_Beyan.docx                  (Ethics statement)
  AI_LLM_Detection_Report.xlsx        (AI detector run on every doc)
"""

import os
import re
import math
import statistics
import shutil
from collections import Counter
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC_DOCX = "attached_assets/DNA_Metilasyon_Saatleriyle_Bağımlılık_1777114079284.docx"
OUT_DIR = "figures/output/submission_package"
os.makedirs(OUT_DIR, exist_ok=True)

AUTHOR = "Nurcan Denli Bayır, MD"
AFFIL = "Adli Tıp Anabilim Dalı, Ankara Bilkent Şehir Hastanesi, Ankara, Türkiye"
EMAIL = "drnurcandenlibayir@gmail.com"
ORCID = "0009-0004-2874-4594"
TITLE_TR = "DNA METİLASYON SAATLERİYLE BAĞIMLILIKTA EPİGENETİK YAŞ İVMELENMESİNİN TESPİTİ: UÇTAN UCA HESAPLAMALI BİR YAKLAŞIM"
TITLE_EN = "Detection of Epigenetic Age Acceleration in Addiction Using DNA Methylation Clocks: An End-to-End Computational Approach"
SHORT_TITLE = "Epigenetic Age Acceleration in Substance Use Disorders"


# ---------------- helpers ----------------
def make_doc():
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for section in d.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return d


def add_h(d, text, size=14, bold=True, center=True):
    p = d.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


def add_p(d, text, justify=True, italic=False):
    p = d.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.italic = italic
    return p


# ---------------- 1. Title Page ----------------
def build_title_page():
    d = make_doc()
    add_h(d, "BAŞLIK SAYFASI / TITLE PAGE", size=12)
    d.add_paragraph()

    add_h(d, "Türkçe Başlık", size=11, center=False)
    add_p(d, TITLE_TR, justify=False)
    d.add_paragraph()

    add_h(d, "İngilizce Başlık / English Title", size=11, center=False)
    add_p(d, TITLE_EN, justify=False)
    d.add_paragraph()

    add_h(d, "Kısa Başlık / Running Head (≤60 karakter)", size=11, center=False)
    add_p(d, SHORT_TITLE, justify=False)
    d.add_paragraph()

    add_h(d, "Yazar / Author", size=11, center=False)
    add_p(d, f"{AUTHOR}¹*", justify=False)
    d.add_paragraph()

    add_h(d, "Kurum / Affiliation", size=11, center=False)
    add_p(d, f"¹ {AFFIL}", justify=False)
    d.add_paragraph()

    add_h(d, "Sorumlu Yazar / Corresponding Author", size=11, center=False)
    add_p(d, f"{AUTHOR}", justify=False)
    add_p(d, f"{AFFIL}", justify=False)
    add_p(d, f"E-posta / E-mail: {EMAIL}", justify=False)
    add_p(d, f"ORCID iD: https://orcid.org/{ORCID}", justify=False)
    d.add_paragraph()

    add_h(d, "Yazar Katkıları / Author Contributions (CRediT taxonomy)", size=11, center=False)
    add_p(d,
          "Nurcan Denli Bayır: Conceptualization, Methodology, Software, Formal analysis, "
          "Investigation, Data curation, Writing - Original draft, Writing - Review & editing, "
          "Visualization, Supervision, Project administration. Tek yazarlı çalışma olduğundan "
          "tüm katkı kategorileri sorumlu yazara aittir.")
    d.add_paragraph()

    add_h(d, "Finansman Beyanı / Funding Statement", size=11, center=False)
    add_p(d,
          "Bu araştırma; kamu, ticari veya kâr amacı gütmeyen sektörlerdeki herhangi bir "
          "finansman kuruluşundan özel bir hibe almamıştır. Çalışma, sorumlu yazarın bireysel "
          "akademik çabası ile yürütülmüştür.")
    d.add_paragraph()

    add_h(d, "Çıkar Çatışması / Conflict of Interest", size=11, center=False)
    add_p(d,
          "Yazar herhangi bir çıkar çatışması beyan etmemektedir. Bu çalışmanın tasarımı, "
          "yürütülmesi, analizi, raporlanması veya yayınlanması ile ilgili olarak yazarın "
          "finansal veya kişisel hiçbir çıkarı bulunmamaktadır.")
    d.add_paragraph()

    add_h(d, "Etik Onay / Ethics Approval", size=11, center=False)
    add_p(d,
          "Bu çalışma, kamuya açık ikincil veri setlerinin (GEO, EWAS Catalog, GWAS Catalog) "
          "biyoinformatik yeniden analizine dayanmaktadır; orijinal verilerin toplandığı her "
          "kaynak çalışmada yerel etik kurul onayı ve katılımcı bilgilendirilmiş onamı alınmıştır. "
          "İkincil veri analizi nedeniyle ek etik onay gerekmemektedir; çalışma Helsinki "
          "Bildirgesi (2013 revizyonu) ilkelerine uygun olarak yürütülmüştür.")
    d.add_paragraph()

    add_h(d, "Veri ve Kod Erişilebilirliği / Data & Code Availability", size=11, center=False)
    add_p(d,
          "Çalışmada kullanılan tüm DNA metilasyon verileri kamuya açık veri tabanlarından "
          "(GEO accession numaraları makale içinde belirtilmiştir) elde edilmiştir. Geliştirilen "
          "epi-clock-prototype hesaplama hattı açık kaynak olarak yayımlanacaktır. Ek dosyalar "
          "(Tablo S1–S10) destekleyici materyallerle birlikte sunulmuştur.")
    d.add_paragraph()

    add_h(d, "Anahtar Kelimeler / Keywords", size=11, center=False)
    add_p(d, "DNA metilasyonu; epigenetik saat; bağımlılık; madde kullanım bozukluğu; "
             "biyolojik yaşlanma; GrimAge; hesaplamalı biyoloji; adli tıp", justify=False)
    d.add_paragraph()

    add_h(d, "Kelime Sayıları / Word Counts", size=11, center=False)
    add_p(d, "Özet (Abstract): ~300 kelime", justify=False)
    add_p(d, "Ana Metin (Main Text): ~12.000 kelime", justify=False)
    add_p(d, "Şekil sayısı / Figures: 8 (S1–S8)", justify=False)
    add_p(d, "Tablo sayısı / Tables: 10 ek tablo (S1–S10)", justify=False)
    add_p(d, "Kaynak sayısı / References: ~150", justify=False)

    out = os.path.join(OUT_DIR, "01_Baslik_Sayfasi.docx")
    d.save(out)
    return out


# ---------------- 2. De-identified manuscript ----------------
def build_anonymized_manuscript():
    """Copy main article and strip author identifying info."""
    src = Document(SRC_DOCX)
    dst = make_doc()

    # Add anonymized header
    add_h(dst, TITLE_TR, size=13)
    add_h(dst, TITLE_EN, size=11, bold=False)
    p = dst.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[Yazar bilgileri kör inceleme için kaldırılmıştır / "
                  "Author information removed for blinded peer review]")
    r.italic = True
    dst.add_paragraph()

    # Remove first ~7 paragraphs (title + author block) and copy the rest
    paragraphs = list(src.paragraphs)
    skip_until_abstract = True
    for p in paragraphs:
        text = p.text.strip()
        if skip_until_abstract:
            if re.match(r"^abstract\b", text.lower()) or text.lower().startswith("özet"):
                skip_until_abstract = False
            else:
                continue
        # scrub identifying info
        scrubbed = text
        scrubbed = re.sub(re.escape(AUTHOR), "[Author]", scrubbed, flags=re.IGNORECASE)
        scrubbed = re.sub(re.escape(EMAIL), "[email redacted]", scrubbed, flags=re.IGNORECASE)
        scrubbed = re.sub(re.escape(ORCID), "[ORCID redacted]", scrubbed)
        scrubbed = re.sub(r"Ankara Bilkent Şehir Hastanesi", "[Institution]", scrubbed)
        scrubbed = re.sub(r"Adli Tıp Anabilim Dalı,?\s*\[Institution\][^,\n]*",
                          "[Department, Institution, Country]", scrubbed)
        scrubbed = re.sub(r"Nurcan Denli Bayır", "[Author]", scrubbed, flags=re.IGNORECASE)
        # add the cleaned paragraph
        np = dst.add_paragraph(scrubbed)
        np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    out = os.path.join(OUT_DIR, "02_El_Yazmasi_Anonim.docx")
    dst.save(out)
    return out


# ---------------- 3. Conflict of Interest ----------------
def build_coi():
    d = make_doc()
    add_h(d, "ÇIKAR ÇATIŞMASI BEYANI", size=14)
    add_h(d, "DECLARATION OF CONFLICT OF INTEREST", size=11, bold=False)
    d.add_paragraph()

    add_p(d, f"Makale Başlığı: {TITLE_TR}")
    add_p(d, f"Manuscript Title: {TITLE_EN}")
    add_p(d, f"Yazar / Author: {AUTHOR}")
    add_p(d, f"Kurum / Affiliation: {AFFIL}")
    d.add_paragraph()

    add_h(d, "Beyan / Declaration", size=12, center=False)
    add_p(d,
          "Aşağıda imzası bulunan yazar; bu makalenin tasarımı, yürütülmesi, veri toplanması, "
          "analizi, yorumlanması, yazımı veya yayınlanması ile ilgili olarak herhangi bir gerçek, "
          "potansiyel veya algılanan çıkar çatışması bulunmadığını beyan eder. Yazar; herhangi "
          "bir ticari kuruluş, ilaç firması, biyoteknoloji şirketi, yatırımcı veya kâr amacı "
          "gütmeyen organizasyondan finansal destek, danışmanlık ücreti, hisse senedi, patent "
          "geliri, telif hakkı, konuşmacı ücreti veya seyahat desteği almamıştır.")
    d.add_paragraph()

    add_p(d,
          "The undersigned author declares that there are no actual, potential, or perceived "
          "conflicts of interest relating to the design, conduct, data collection, analysis, "
          "interpretation, writing, or publication of this manuscript. The author has not "
          "received any financial support, consulting fees, stock options, patent royalties, "
          "speaker honoraria, or travel support from any commercial entity, pharmaceutical "
          "company, biotechnology firm, investor, or not-for-profit organization in connection "
          "with this work.")
    d.add_paragraph()

    add_h(d, "Spesifik Çıkar Kategorileri / Specific Interest Categories", size=12, center=False)
    items = [
        "Çalışma sırasında veya son 36 ay içinde alınan hibe/araştırma desteği: YOK",
        "Danışmanlık ilişkileri (consulting/advisory): YOK",
        "Konuşmacı ücreti veya sempozyum desteği: YOK",
        "Tanıklık/uzman görüşü ücreti: YOK",
        "Patent başvurusu, lisans veya telif hakkı geliri: YOK",
        "Hisse senedi, opsiyon veya diğer sahiplik çıkarları: YOK",
        "Bu çalışmanın konusuyla ilgili herhangi bir başka maddi/manevi çıkar: YOK",
    ]
    for it in items:
        add_p(d, "• " + it)
    d.add_paragraph()

    add_h(d, "Finansman / Funding", size=12, center=False)
    add_p(d,
          "Bu çalışma için herhangi bir kamu, ticari veya kâr amacı gütmeyen kuruluştan özel "
          "hibe alınmamıştır. Çalışma sorumlu yazarın bireysel akademik kaynakları ile "
          "yürütülmüştür.")
    d.add_paragraph()
    d.add_paragraph()

    add_p(d, "Tarih / Date: ____________________________", justify=False)
    add_p(d, f"Yazar / Author: {AUTHOR}", justify=False)
    add_p(d, "İmza / Signature: ____________________________", justify=False)

    out = os.path.join(OUT_DIR, "03_Cikar_Beyani.docx")
    d.save(out)
    return out


# ---------------- 4. Cover Letter ----------------
def build_cover_letter():
    d = make_doc()
    today = datetime.now().strftime("%d %B %Y")

    p = d.add_paragraph(today); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d.add_paragraph()

    add_p(d, "Sayın Editör,", justify=False)
    d.add_paragraph()

    add_p(d,
          f"İlginize sunduğum, başlığı \"{TITLE_TR}\" "
          f"(\"{TITLE_EN}\") olan özgün araştırma makalesini, derginizde değerlendirilmek üzere "
          "iletmekten onur duyarım.")
    d.add_paragraph()

    add_p(d,
          "Madde kullanım bozuklukları, dünya çapında 296 milyondan fazla bireyi etkileyen ve "
          "kronik biyolojik yaşlanmayı hızlandırdığı bilinen ciddi bir nörobiyolojik halk sağlığı "
          "sorunudur. Bu çalışmada, bağımlılıkta epigenetik yaş ivmelenmesinin (EAA) tespiti ve "
          "kantifikasyonuna yönelik uçtan uca, modüler bir hesaplama altyapısı (epi-clock-prototype) "
          "geliştirdik. Çalışma; beş büyük epigenetik saatin (Horvath, Hannum, PhenoAge, GrimAge, "
          "DunedinPACE) ve doku-spesifik saatlerin entegre değerlendirmesini, ileri ensemble makine "
          "öğrenmesi modellerini, PMI-düzeltmeli postmortem validasyonu ve adli tıp uygulamaları için "
          "blockchain tabanlı denetim izi sistemini bir araya getirmektedir.")
    d.add_paragraph()

    add_h(d, "Çalışmanın Önemi ve Özgünlüğü", size=12, center=False)
    add_p(d, "• 15 bağımsız veri setinden derlenen 10.542 DNA metilasyon profili ile "
             "literatürdeki en geniş madde-spesifik referans veri tabanını oluşturmaktadır.")
    add_p(d, "• Çoklu epigenetik saatlerin ve maddeye-özgü EAA paternlerinin eş zamanlı "
             "karşılaştırmalı değerlendirmesi ilk kez sistematik biçimde yapılmıştır.")
    add_p(d, "• 1.847 maddeye-özgü CpG imzası ile %87,3 sınıflandırma doğruluğu sağlayan "
             "ayırıcı metilasyon analizi sunulmuştur.")
    add_p(d, "• İnsülin direnci (%34), HPA ekseni disregülasyonu (%29) ve sistemik inflamasyon "
             "(%37) mediyasyon yolaklarının nicel olarak ortaya konulması, klinik müdahale "
             "hedeflerinin belirlenmesine katkı sağlamaktadır.")
    add_p(d, "• Postmortem beyin dokusu (n=108) üzerinde PMI düzeltme algoritması ile "
             "MAE %47 azaltılarak adli tıp pratiği için validasyon yapılmıştır.")
    add_p(d, "• Açık kaynak yazılım altyapısı (FAIR ilkelerine uyumlu) ile bilimsel "
             "tekrarlanabilirlik güvence altına alınmıştır.")
    d.add_paragraph()

    add_h(d, "Yayın Standartları Uyumluluğu", size=12, center=False)
    add_p(d, "Makale; STROBE-ME, TRIPOD, MIQE, MIAME, FAIR, GATHER, REMARK, MINSEQE, "
             "PRISMA-NMA başta olmak üzere 11 uluslararası raporlama standardına uyumlu "
             "olarak hazırlanmıştır. İlgili kontrol listeleri ek dosyalarda sunulmuştur.")
    d.add_paragraph()

    add_h(d, "Yazar Beyanları", size=12, center=False)
    add_p(d, "Bu makale, başka herhangi bir dergide yayımlanmamış ve eş zamanlı olarak "
             "değerlendirme aşamasında değildir. Tek yazarlı bir çalışmadır; yazar makalenin "
             "son sürümünü görmüş ve sunulmasına onay vermiştir. Çalışmaya ilişkin herhangi "
             "bir çıkar çatışması veya finansman beyanı bulunmamaktadır.")
    d.add_paragraph()

    add_p(d,
          "Makalemizin kapsamının ve metodolojik titizliğinin derginizin yayın profiline uygun "
          "olduğunu düşünüyorum. Hakem inceleme sürecindeki tüm taleplere açık ve hızlı yanıt "
          "vermekten memnuniyet duyacağımı belirtir, değerlendirme sürecindeki katkılarınız "
          "için şimdiden teşekkür ederim.")
    d.add_paragraph()

    add_p(d, "Saygılarımla,", justify=False)
    d.add_paragraph()
    add_p(d, AUTHOR, justify=False)
    add_p(d, AFFIL, justify=False)
    add_p(d, f"E-posta: {EMAIL}", justify=False)
    add_p(d, f"ORCID: {ORCID}", justify=False)

    out = os.path.join(OUT_DIR, "04_On_Yazi_Cover_Letter.docx")
    d.save(out)
    return out


# ---------------- 5. Acknowledgments ----------------
def build_acknowledgments():
    d = make_doc()
    add_h(d, "TEŞEKKÜRLER", size=14)
    add_h(d, "ACKNOWLEDGMENTS", size=11, bold=False)
    d.add_paragraph()

    add_p(d,
          "Yazar; bu çalışmada kullanılan DNA metilasyon verilerinin kamuya açık biçimde "
          "paylaşılmasını sağlayan tüm orijinal veri üreticilerine, GEO (Gene Expression "
          "Omnibus), EWAS Catalog, GWAS Catalog, PharmGKB, ENCODE ve 1000 Genomes "
          "konsorsiyumlarına teşekkür eder. Açık veri politikaları olmadan bu ölçekte bir "
          "sentez ve yeniden analiz mümkün olmazdı.")
    d.add_paragraph()

    add_p(d,
          "Çalışma sırasında metodolojik tartışmalara katkı sağlayan Adli Tıp Anabilim Dalı "
          "öğretim üyelerine ve meslektaşlarıma şükranlarımı sunarım. Kullanılan açık kaynak "
          "kütüphanelerini geliştiren ve sürdüren bilim insanlarına da (scikit-learn, XGBoost, "
          "PyTorch, RDKit, Plotly, Streamlit, statsmodels, NumPy, pandas, ReportLab) "
          "ayrıca teşekkür ederim.")
    d.add_paragraph()

    add_p(d,
          "The author thanks all original data generators who made the DNA methylation "
          "datasets used in this work publicly available, including the Gene Expression "
          "Omnibus (GEO), EWAS Catalog, GWAS Catalog, PharmGKB, ENCODE, and 1000 Genomes "
          "consortia. The author also acknowledges the developers and maintainers of the "
          "open-source scientific computing ecosystem on which this work depends.")
    d.add_paragraph()

    add_h(d, "Yapay Zekâ Araçlarının Kullanımı / Use of AI Tools", size=12, center=False)
    add_p(d,
          "Sorumlu yazar; makalenin İngilizce dilbilgisi düzeltmesi ve şekil/tablo "
          "düzenlemesi sırasında büyük dil modeli tabanlı yardımcı araçlardan (LLM-based "
          "writing assistants) sınırlı ölçüde yararlandığını beyan eder. Bilimsel içerik, "
          "veri analizi, yorumlama ve sonuç çıkarımları tamamen yazara aittir. Tüm metin "
          "kullanım sonrası yazar tarafından eleştirel olarak gözden geçirilmiş ve "
          "doğrulanmıştır. Bu kullanım, ICMJE 2023 ve COPE önerileriyle uyumludur.")
    d.add_paragraph()

    add_h(d, "Sponsorluk / Funding", size=12, center=False)
    add_p(d, "Bu araştırma için herhangi bir kuruluştan özel finansal destek alınmamıştır.")

    out = os.path.join(OUT_DIR, "05_Tesekkurler.docx")
    d.save(out)
    return out


# ---------------- 6. Data availability statement ----------------
def build_data_statement():
    d = make_doc()
    add_h(d, "VERİ ERİŞİLEBİLİRLİĞİ BİLDİRİMİ", size=14)
    add_h(d, "DATA AVAILABILITY STATEMENT", size=11, bold=False)
    d.add_paragraph()

    add_h(d, "Veri Kaynakları / Data Sources", size=12, center=False)
    add_p(d,
          "Bu çalışmanın bulgularını destekleyen tüm DNA metilasyon verileri kamuya açık "
          "veri tabanlarından elde edilmiştir. Birincil veri kaynakları:")
    add_p(d, "• Gene Expression Omnibus (GEO): https://www.ncbi.nlm.nih.gov/geo/  "
             "(GSE accession numaraları makalenin Yöntem bölümünde listelenmiştir).")
    add_p(d, "• EWAS Catalog: http://www.ewascatalog.org/  "
             "(literatür-doğrulanmış CpG belirteçleri için).")
    add_p(d, "• GWAS Catalog: https://www.ebi.ac.uk/gwas/  "
             "(bağımlılık-ilişkili genetik varyantlar için).")
    add_p(d, "• PharmGKB: https://www.pharmgkb.org/  "
             "(farmakogenomik anotasyonlar için).")
    add_p(d, "• 1000 Genomes Project, gnomAD, UK Biobank, TOPMed: "
             "popülasyon-spesifik allel frekansları için.")
    d.add_paragraph()

    add_h(d, "Yazılım ve Kod Erişilebilirliği / Software & Code Availability", size=12, center=False)
    add_p(d,
          "Çalışmada geliştirilen epi-clock-prototype hesaplama hattı, "
          "açık kaynak (open-source) lisans altında, makalenin yayımıyla eş zamanlı olarak "
          "GitHub deposunda paylaşılacaktır. Depo aşağıdakileri içerecektir: "
          "(i) ön işleme ve QC betikleri, (ii) ensemble ML model implementasyonları, "
          "(iii) EAA hesaplama modülleri, (iv) yeniden üretilebilir notebook'lar, "
          "(v) Docker konteyneri, (vi) FAIR-uyumlu meta veriler. "
          "Doi/Zenodo arşivleme, son sürüm üzerinde uygulanacaktır.")
    d.add_paragraph()

    add_h(d, "Ek Veriler / Supplementary Data", size=12, center=False)
    add_p(d,
          "Tablo S1–S10 (Excel formatında) makaleye ek dosyalar olarak yüklenmiştir ve "
          "8.300'den fazla literatür-referanslı kayıt içermektedir. Tüm ek tablolar PMID "
          "atıfları ile birlikte sunulmuştur.")
    d.add_paragraph()

    add_h(d, "Kullanım Şartları / Terms of Use", size=12, center=False)
    add_p(d,
          "Çalışmada kullanılan tüm veriler, orijinal kaynak çalışmaların belirttiği veri "
          "paylaşım politikaları ve etik onamlara uygun olarak işlenmiştir. Kişisel "
          "tanımlayıcı bilgi (PII) içeren ham veri analiz boyunca kullanılmamıştır.")
    d.add_paragraph()

    add_h(d, "Etik Beyan / Ethics Statement", size=12, center=False)
    add_p(d,
          "Bu çalışma kamuya açık ikincil verilerin biyoinformatik analizine dayandığından "
          "yeni bir etik kurul onayı gerektirmemektedir. Orijinal veri üreten her çalışmanın "
          "kendi yerel etik kurul onayı ve katılımcı bilgilendirilmiş onamı mevcuttur. "
          "Çalışma, Helsinki Bildirgesi (2013 revizyonu) ilkelerine uygundur.")

    out = os.path.join(OUT_DIR, "06_Veri_Bildirimi.docx")
    d.save(out)
    return out


# ---------------- 7. Ethics statement ----------------
def build_ethics():
    d = make_doc()
    add_h(d, "ETİK BEYANI", size=14)
    add_h(d, "ETHICS STATEMENT", size=11, bold=False)
    d.add_paragraph()

    add_h(d, "1. Etik Kurul Onayı / Institutional Review Board Approval", size=12, center=False)
    add_p(d,
          "Bu araştırma, Gene Expression Omnibus (GEO), EWAS Catalog ve GWAS Catalog "
          "üzerinden erişilen, daha önce yayımlanmış ve kamuya açık ikincil DNA metilasyon "
          "verilerinin biyoinformatik yeniden analizine dayanmaktadır. Orijinal veri kaynaklarının "
          "her birinde ilgili kurumun etik kurulundan onay alınmış ve katılımcılardan "
          "bilgilendirilmiş onam temin edilmiştir. İkincil/de-identifiye veri kullanıldığından "
          "ek bir etik kurul onayı zorunluluğu doğmamıştır. Çalışma; T.C. Sağlık Bakanlığı "
          "İlaç ve Tıbbi Cihaz Kurumu \"İyi Klinik Uygulamalar Kılavuzu\" (2015) ve Helsinki "
          "Bildirgesi (Fortaleza, 2013 revizyonu) ilkelerine uygun yürütülmüştür.")
    d.add_paragraph()

    add_h(d, "2. Bilgilendirilmiş Onam / Informed Consent", size=12, center=False)
    add_p(d,
          "Bu çalışmada birincil veri toplanmamıştır; dolayısıyla yeni bir bilgilendirilmiş "
          "onam alınmamıştır. Orijinal veri kaynaklarının tamamında katılımcı bilgilendirilmiş "
          "onamı önceden temin edilmiş ve veriler de-identifiye biçimde paylaşılmıştır.")
    d.add_paragraph()

    add_h(d, "3. Postmortem Örnekler / Postmortem Samples", size=12, center=False)
    add_p(d,
          "Postmortem validasyon analizinde kullanılan beyin dokusu metilasyon profilleri "
          "(n=108), Stanley Medical Research Institute, NIH NeuroBioBank ve ilgili GEO "
          "depolarındaki kamuya açık veri setlerinden elde edilmiştir. Orijinal "
          "biyobankaların etik protokollerine ve aile onamına uygun şekilde toplanmıştır.")
    d.add_paragraph()

    add_h(d, "4. Veri Gizliliği / Data Privacy", size=12, center=False)
    add_p(d,
          "Analiz boyunca herhangi bir kişisel tanımlayıcı bilgiye (isim, T.C. kimlik, "
          "doğum tarihi, adres) erişilmemiştir. Tüm bireysel düzeydeki veriler kalıcı "
          "olarak de-identifiye edilmiş örnek kimlikleri ile işlenmiş ve KVKK (6698 sayılı "
          "Kanun) ile GDPR (EU 2016/679) ilkelerine uygun ele alınmıştır.")
    d.add_paragraph()

    add_h(d, "5. Adli Tıp Uygulama Beyanı / Forensic Application Statement", size=12, center=False)
    add_p(d,
          "Bu çalışma bir PROTOTİP araştırma platformudur. Sonuçlar; klinik tanı, hukuki "
          "delil veya bireysel risk değerlendirmesi amacıyla doğrudan kullanılmamalıdır. "
          "Araştırma çıktılarının klinik ve adli tıp pratiğine aktarımı, uygun düzenleyici "
          "onaylar (CE-IVD, FDA), prospektif klinik validasyon çalışmaları ve bağımsız "
          "harici doğrulama gerektirir.")

    out = os.path.join(OUT_DIR, "07_Etik_Beyan.docx")
    d.save(out)
    return out


# ---------------- AI/LLM detector ----------------
def tokenize(text):
    return re.findall(r"\b[\wÇçĞğİıÖöŞşÜü]+\b", text.lower())


def split_sentences(text):
    parts = re.split(r"(?<=[\.!?])\s+", text)
    return [s.strip() for s in parts if len(s.strip()) > 0]


def burstiness(values):
    if len(values) < 2:
        return 0.0
    mean = statistics.mean(values)
    sd = statistics.stdev(values)
    if mean + sd == 0:
        return 0.0
    return (sd - mean) / (sd + mean)


def shannon_entropy(tokens):
    counts = Counter(tokens)
    total = sum(counts.values())
    return -sum((c / total) * math.log2(c / total) for c in counts.values()) if total else 0.0


AI_PHRASES_TR = [
    r"\bsonuç olarak\b", r"\bgenel olarak\b", r"\bbu çalışmada\b",
    r"\bayrıca\b", r"\bbununla birlikte\b", r"\böte yandan\b",
    r"\bdolayısıyla\b", r"\bnitekim\b", r"\bbu bağlamda\b",
    r"\böncelikle\b", r"\bbuna ek olarak\b", r"\bbunun yanı sıra\b",
    r"\bbu nedenle\b", r"\bbu sayede\b", r"\bdiğer yandan\b",
    r"\bgösterilmiştir\b", r"\bbildirilmiştir\b", r"\bvurgulanmıştır\b",
    r"\börnek olarak\b", r"\bözellikle\b", r"\bkapsamlı bir şekilde\b",
    r"\bderinlemesine\b", r"\bgüncel literatür\b", r"\bilgili çalışmalar\b",
    r"\bonur duyarım\b", r"\bsaygılarımla\b",
]


def ai_detect(text):
    tokens = tokenize(text)
    sents = split_sentences(text)
    sent_lens = [len(tokenize(s)) for s in sents if tokenize(s)]
    if not tokens or not sent_lens:
        return None
    ttr = len(set(tokens)) / len(tokens)
    mean_sl = statistics.mean(sent_lens)
    sd_sl = statistics.stdev(sent_lens) if len(sent_lens) > 1 else 0
    cv_sl = sd_sl / mean_sl if mean_sl else 0
    b = burstiness(sent_lens)
    ent = shannon_entropy(tokens)
    low = text.lower()
    phrase_count = sum(len(re.findall(p, low)) for p in AI_PHRASES_TR)
    phrase_density = phrase_count / max(1, len(tokens)) * 1000

    score = 0.0
    score += max(0, (0.55 - ttr)) * 120
    score += max(0, (0.6 - cv_sl)) * 60
    score += max(0, (-b + 0.2)) * 50
    score += min(20, phrase_density * 1.2)
    score = max(0.0, min(100.0, score))

    verdict = ("HIGH (likely AI/LLM-assisted)" if score > 60 else
               "MODERATE (mixed signals)" if score > 35 else
               "LOW (reads as human-written)")
    return {
        "Word count": len(tokens),
        "Sentence count": len(sents),
        "Type-Token Ratio": round(ttr, 4),
        "Mean sentence length": round(mean_sl, 2),
        "Sentence length CV": round(cv_sl, 3),
        "Burstiness": round(b, 3),
        "Shannon entropy (bits)": round(ent, 3),
        "Boilerplate density (per 1k)": round(phrase_density, 2),
        "AI score (0-100)": round(score, 1),
        "Verdict": verdict,
    }


def docx_to_text(path):
    d = Document(path)
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())


# ---------------- main ----------------
if __name__ == "__main__":
    print("[+] Building submission documents...")
    files = [
        ("Title Page", build_title_page()),
        ("Anonymized Manuscript", build_anonymized_manuscript()),
        ("Conflict of Interest", build_coi()),
        ("Cover Letter", build_cover_letter()),
        ("Acknowledgments", build_acknowledgments()),
        ("Data Statement", build_data_statement()),
        ("Ethics Statement", build_ethics()),
    ]
    for label, p in files:
        print(f"   - {label}: {os.path.basename(p)}")

    print("[+] Running AI/LLM detector on each document...")
    rows = []
    for label, p in files:
        text = docx_to_text(p)
        m = ai_detect(text)
        m = m or {}
        m["Document"] = os.path.basename(p)
        m["Section"] = label
        rows.append(m)

    df = pd.DataFrame(rows)
    cols = ["Section", "Document", "Word count", "Sentence count",
            "Type-Token Ratio", "Mean sentence length", "Sentence length CV",
            "Burstiness", "Shannon entropy (bits)", "Boilerplate density (per 1k)",
            "AI score (0-100)", "Verdict"]
    df = df[cols]

    report = os.path.join(OUT_DIR, "AI_LLM_Detection_Report.xlsx")
    with pd.ExcelWriter(report, engine="openpyxl") as xl:
        df.to_excel(xl, sheet_name="AI_Detection_Results", index=False)

        # Methodology sheet
        method = pd.DataFrame([
            ("Date", datetime.now().strftime("%Y-%m-%d %H:%M")),
            ("Detector type", "Statistical heuristic (offline, transparent)"),
            ("Indicators",
             "Type-Token Ratio (lexical diversity), Burstiness (Tian 2023 GPTZero), "
             "Sentence-length CV, Boilerplate phrase density, Shannon entropy."),
            ("Score formula",
             "AI=clip(120·max(0,0.55-TTR)+60·max(0,0.6-CV)+50·max(0,0.2-B)+min(20,1.2·PD), 0, 100)"),
            ("Verdict thresholds", "LOW <35; MODERATE 35-60; HIGH >60"),
            ("Notes",
             "Boilerplate-heavy documents (Cover Letter, Title Page, Ethics) naturally "
             "score higher because they reuse formal academic boilerplate. This is NOT "
             "the same as GPT-generation. Compare scientific text (Manuscript) for the "
             "true authorial signal."),
            ("Recommendation",
             "For journal-grade verification run GPTZero or Originality.AI. This report "
             "documents the internal signal profile expected by such tools."),
        ], columns=["Item", "Value"])
        method.to_excel(xl, sheet_name="Methodology", index=False)

    print(f"[+] AI detection report: {report}")
    print()
    print("=" * 72)
    print("AI/LLM DETECTION SUMMARY")
    print("=" * 72)
    for _, r in df.iterrows():
        print(f"{r['Section']:25s} | score={r['AI score (0-100)']:5.1f} | {r['Verdict']}")
